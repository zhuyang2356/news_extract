from docx import Document
from zipfile import ZipFile
from bs4 import BeautifulSoup

def extract_docx(path):
    try:
        doc = Document(path)
        # print(doc.paragraphs)
        doc_text_list=[]
        for paragraph in doc.paragraphs:
        #     print(paragraph.text)
            if len(paragraph.text)>=5:
                doc_text_list.append(paragraph.text)
        return doc_text_list
    except Exception as e:
        doc_text_list=[]
        doc=ZipFile(path)
        xml=doc.read("word/document.xml")
        wordObj = BeautifulSoup(xml.decode("utf-8"))
        texts = wordObj.findAll("w:t")
        for text in texts:
            doc_text_list.append(text.text)
        return doc_text_list


